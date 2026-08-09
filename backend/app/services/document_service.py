import os
import uuid
import aiofiles

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.config import get_settings

settings = get_settings()

ALLOWED_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}

MAX_FILE_SIZE = settings.MAX_FILE_SIZE_MB * 1024 * 1024


def validate_file(filename: str, content_type: str, file_size: int) -> str | None:
    if content_type not in ALLOWED_TYPES:
        return f"Unsupported file type: {content_type}. Only PDF and DOCX are supported."
    if file_size > MAX_FILE_SIZE:
        return f"File size exceeds the maximum allowed size of {settings.MAX_FILE_SIZE_MB} MB."
    return None


async def save_upload(content: bytes, original_filename: str, tenant_id: int) -> dict:
    ext = os.path.splitext(original_filename)[1] or ".bin"
    safe_filename = f"{uuid.uuid4()}{ext}"
    tenant_dir = os.path.join(settings.UPLOAD_DIR, str(tenant_id))
    os.makedirs(tenant_dir, exist_ok=True)
    storage_path = os.path.join(tenant_dir, safe_filename)

    async with aiofiles.open(storage_path, "wb") as f:
        await f.write(content)

    return {
        "storage_path": storage_path,
        "file_size": len(content),
        "filename": safe_filename,
    }


def extract_text_from_pdf(file_path: str) -> list[dict]:
    pages = []
    doc = fitz.open(file_path)
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text()
        if text.strip():
            pages.append({"page_number": page_num, "content": text.strip()})
    doc.close()
    return pages


def extract_text_from_docx(file_path: str) -> list[dict]:
    pages = []
    doc = DocxDocument(file_path)

    sections = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        heading = ""
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            heading = f"[{para.style.name}] "
        sections.append(f"{heading}{text}")

    body_text = "\n".join(sections)

    for table in doc.tables:
        table_text = ""
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_text += row_text + "\n"
        if table_text.strip():
            body_text += "\n\n[TABLE]\n" + table_text.strip()

    if body_text.strip():
        pages.append({"page_number": 1, "content": body_text.strip()})

    return pages
